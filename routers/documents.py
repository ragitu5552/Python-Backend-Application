from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
from services.embedding import embedding_service
from typing import Optional
from models import Document, DocumentChunk
from database import get_db
from schemas import DocumentCreate, DocumentResponse, DocumentListResponse, DocumentUpdate
import os
import uuid
import logging

router = APIRouter(prefix="/documents", tags=["documents"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

logger = logging.getLogger(__name__)

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    generate_embeddings: Optional[bool] = True
):
    try:
        # Validate file is not empty
        if not file.filename:
            logger.error("No file uploaded")
            raise HTTPException(status_code=400, detail="No file uploaded")

        # Validate file size (optional)
        file.file.seek(0, os.SEEK_END)
        file_size = file.file.tell()
        file.file.seek(0)  # Reset file pointer

        if file_size > 10_000_000:  # 10MB limit
            logger.error(f"File too large: {file_size} bytes")
            raise HTTPException(status_code=413, detail="File too large")

        # Save file locally
        file_ext = os.path.splitext(file.filename)[1].lower()
        file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}{file_ext}")

        # Ensure uploads directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # Save file content
        try:
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
        except IOError as io_err:
            logger.error(f"File write error: {io_err}")
            raise HTTPException(status_code=500, detail=f"Failed to save file: {io_err}")

        # Extract text (with improved text extraction)
        content_text = ""
        try:
            if file_ext == ".txt":
                with open(file_path, "r", encoding='utf-8') as f:
                    content_text = f.read()
            elif file_ext == ".pdf":
                # Add PDF text extraction (using PyPDF2 or similar)
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    content_text = " ".join(page.extract_text() for page in pdf_reader.pages)
            elif file_ext in [".docx", ".doc"]:
                # Add Word document text extraction (using python-docx or similar)
                import docx
                doc = docx.Document(file_path)
                content_text = " ".join(para.text for para in doc.paragraphs)
        except Exception as read_err:
            logger.warning(f"Could not read file content: {read_err}")
            raise HTTPException(status_code=400, detail=f"Unable to extract text from file: {read_err}")

        # Create document record
        document = Document(
            title=file.filename,
            content=content_text,
            file_path=file_path,
            is_active=True
        )

        # Generate document-level embedding
        document_embedding = None
        if generate_embeddings and content_text:
            try:
                document_embedding = await embedding_service.generate_embeddings(content_text)
                document.embedding = document_embedding
            except Exception as embed_err:
                logger.error(f"Document-level embedding generation error: {embed_err}")

        # Chunk the document and generate chunk embeddings
        chunks = []
        if generate_embeddings and content_text:
            try:
                # Use the chunk_and_embed method from previous implementation
                chunk_texts, chunk_embeddings = await embedding_service.chunk_and_embed(content_text)
                
                # Create DocumentChunk instances
                for text, embedding in zip(chunk_texts, chunk_embeddings):
                    chunk = DocumentChunk(
                        document=document,
                        text=text,
                        embedding=embedding,
                        metadata={
                            "source_file": file.filename,
                            "total_document_length": len(content_text)
                        }
                    )
                    chunks.append(chunk)
            except Exception as chunk_err:
                logger.error(f"Document chunking error: {chunk_err}")

        # Add document and chunks to the session
        db.add(document)
        if chunks:
            db.add_all(chunks)

        await db.commit()
        await db.refresh(document)
        
        # Refresh chunks to ensure they have IDs
        if chunks:
            for chunk in chunks:
                await db.refresh(chunk)

        logger.info(f"Successfully uploaded document: {file.filename}")
        
        # Return the document with some metadata about chunking
        return {
            **document.__dict__,
            "num_chunks": len(chunks),
            "total_document_length": len(content_text)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error uploading document: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@router.get("/", response_model=list[DocumentListResponse])
async def list_documents(
    db: AsyncSession = Depends(get_db),
    skip: int = 0,
    limit: int = 100
):
    result = await db.execute(
        select(Document).offset(skip).limit(limit)
    )
    documents = result.scalars().all()
    return documents

@router.get("/active", response_model=list[DocumentListResponse])
async def list_active_documents(
    db: AsyncSession = Depends(get_db),
    skip: int = 0,
    limit: int = 100
):
    result = await db.execute(
        select(Document)
        .where(Document.is_active == True)
        .offset(skip)
        .limit(limit)
    )
    documents = result.scalars().all()
    return documents

@router.put("/{doc_id}/activate", response_model=DocumentResponse)
async def activate_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db)
):
    await db.execute(
        update(Document)
        .where(Document.id == doc_id)
        .values(is_active=True)
    )
    await db.commit()
    
    result = await db.execute(
        select(Document).where(Document.id == doc_id)
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return document

@router.put("/{doc_id}/deactivate", response_model=DocumentResponse)
async def deactivate_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db)
):
    await db.execute(
        update(Document)
        .where(Document.id == doc_id)
        .values(is_active=False)
    )
    await db.commit()
    
    result = await db.execute(
        select(Document).where(Document.id == doc_id)
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return document

@router.get("/active-state")
async def get_active_documents(db: AsyncSession = Depends(get_db)):
    """Diagnostic endpoint to check currently active documents"""
    result = await db.execute(
        select(Document.id, Document.title, Document.is_active)
        .order_by(Document.is_active.desc(), Document.title)
    )
    return result.all()